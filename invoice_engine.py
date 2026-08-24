"""
Core logic for generating distributor commission-payout TAX INVOICEs
(Word .docx) from an uploaded Excel sheet, following the same layout as
the sample "CP Payout" invoice, and applying these GST rules:

  - No GST No on file (blank / 0)               -> no tax lines, invoice
                                                    amount = Total as-is.
  - GST No present, NOT starting with "33"       -> the GST column value
                                                    is charged in full as
                                                    IGST.
  - GST No present, starting with "33"           -> the GST column value
                                                    is split in half
                                                    between SGST and CGST.

IMPORTANT (updated per user): the "Total" column in the Excel sheet is
now the BASE (pre-tax) commission amount, shown as-is in the
Particulars line. The tax amount is no longer derived mathematically -
it is read directly from the sheet's "GST" column (a rupee value) and
added on top:

    base (Particulars amount) = Total
    tax  (IGST, or SGST+CGST)  = GST column value (split if GST No
                                  starts with "33")
    grand total                = base + tax
"""

import math
import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Fixed "Myalternates" details - identical on every invoice (from sample doc)
# ---------------------------------------------------------------------------
PAYEE_NAME = "Myalternates Financial Services Private Limited"
PAYEE_ADDRESS_LINES = [
    "1B, 2nd Floor, Wellingdon Estaes,",
    "Ethiraj Salai, Egmore",
    "Chennai\u2013 600 008.",
]
PAYEE_GSTIN = "33AAJCP2335B1ZX"
PLACE_OF_SUPPLY = "TAMIL NADU"
SAC_CODE = "997153"

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "myalternates_logo.png")

_ONES = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT",
          "NINE", "TEN", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN",
          "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
_TENS = ["", "", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY",
          "EIGHTY", "NINETY"]


def _two_digits(n):
    if n < 20:
        return _ONES[n]
    return (_TENS[n // 10] + (" " + _ONES[n % 10] if n % 10 else "")).strip()


def _three_digits(n):
    if n >= 100:
        rest = n % 100
        return (_ONES[n // 100] + " HUNDRED" +
                (" " + _two_digits(rest) if rest else "")).strip()
    return _two_digits(n)


def number_to_indian_words(amount):
    """Convert a non-negative integer rupee amount to words, Indian
    numbering system (crore / lakh / thousand)."""
    n = int(round(amount))
    if n == 0:
        return "ZERO"

    crore, n = divmod(n, 10_000_000)
    lakh, n = divmod(n, 100_000)
    thousand, n = divmod(n, 1_000)
    hundred = n

    parts = []
    if crore:
        parts.append(_three_digits(crore) + " CRORE")
    if lakh:
        parts.append(_three_digits(lakh) + " LAKH")
    if thousand:
        parts.append(_three_digits(thousand) + " THOUSAND")
    if hundred:
        parts.append(_three_digits(hundred))
    return " ".join(parts)


def is_blank_gst(value):
    """True if a GST-number cell should be treated as 'no GST' (blank, 0,
    '0', NaN, None, etc.)."""
    if value is None:
        return True
    if isinstance(value, (int, float)):
        try:
            return float(value) == 0 or math.isnan(float(value))
        except (TypeError, ValueError):
            return True
    s = str(value).strip()
    return s == "" or s == "0" or s.lower() == "nan"


def _safe_amount(value):
    """Coerce a cell value to a non-negative float rupee amount, treating
    blank/NaN/None as 0."""
    if value is None:
        return 0.0
    try:
        f = float(value)
        if math.isnan(f):
            return 0.0
        return f
    except (TypeError, ValueError):
        return 0.0


def compute_gst(total_amount, gst_no, gst_amount=0):
    """Given the Excel 'Total' (base, pre-tax commission amount), the GST
    number, and the Excel 'GST' column (the tax rupee value), return a
    dict with base amount, igst, sgst, cgst, the combined tax_total and
    the reconciled grand total, all rounded to whole rupees.

    Rules:
      - blank/0 GST No -> no tax at all, regardless of the GST column.
      - GST No starts with "33" -> GST column value split into SGST +
        CGST halves (rounded so the two halves still sum exactly to the
        original value).
      - GST No present, otherwise -> GST column value charged in full as
        IGST.
    """
    base = round(_safe_amount(total_amount))

    if is_blank_gst(gst_no):
        return {
            "gst_applicable": False,
            "base": base,
            "igst": 0,
            "sgst": 0,
            "cgst": 0,
            "tax_total": 0,
            "total": base,
        }

    tax_total = round(_safe_amount(gst_amount))
    gst_clean = str(gst_no).strip()

    if gst_clean.startswith("33"):
        sgst = round(tax_total / 2)
        cgst = tax_total - sgst  # keeps sgst + cgst exactly == tax_total
        return {
            "gst_applicable": True,
            "base": base,
            "igst": 0,
            "sgst": sgst,
            "cgst": cgst,
            "tax_total": tax_total,
            "total": base + tax_total,
        }
    else:
        return {
            "gst_applicable": True,
            "base": base,
            "igst": tax_total,
            "sgst": 0,
            "cgst": 0,
            "tax_total": tax_total,
            "total": base + tax_total,
        }


def fmt_rupees(n):
    """1234567 -> '12,34,567' using the Indian digit grouping."""
    n = int(n)
    s = str(abs(n))
    if len(s) <= 3:
        grouped = s
    else:
        last3 = s[-3:]
        rest = s[:-3]
        parts = []
        while len(rest) > 2:
            parts.insert(0, rest[-2:])
            rest = rest[:-2]
        if rest:
            parts.insert(0, rest)
        grouped = ",".join(parts) + "," + last3
    return ("-" if n < 0 else "") + grouped


def _no_border_table(table):
    """Strip all borders from a table (used for label/value layouts that
    should look like plain aligned text, not a grid)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _shade_cell(cell, hex_color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold=False, align=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


def _set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """Set cell padding in twips (1/20 pt) so table rows look spacious."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def _set_row_height(row, twips=460):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def _tight(paragraph, space_after=2, space_before=0):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = 1.0
    return paragraph


def _add_logo_header(doc):
    """Place the Myalternates logo top-right of every page."""
    if not os.path.isfile(LOGO_PATH):
        return
    header = doc.sections[0].header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(3.4))


def build_invoice_docx(row, month_period, invoice_date, ref_prefix, out_path):
    """Build a single distributor TAX INVOICE .docx.

    row: dict with keys partner_code, partner_name, total, gst_no,
         gst_amount, bank_name, account_number, ifsc
    month_period: e.g. 'Apr-2023' (shown in Ref No. and Particulars line)
    invoice_date: string, e.g. '30-Apr-2023'
    ref_prefix: text before the '/', normally the partner code
    """
    gst = compute_gst(row["total"], row["gst_no"], row.get("gst_amount", 0))

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    _add_logo_header(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(title, space_after=8)
    r = title.add_run("TAX INVOICE")
    r.bold = True
    r.font.size = Pt(15)

    _tight(doc.add_paragraph(f"Ref. No: {ref_prefix}/{month_period}"))
    _tight(doc.add_paragraph(f"Date: {invoice_date}"), space_after=10)

    _tight(doc.add_paragraph("From,"))
    p = doc.add_paragraph()
    _tight(p)
    r = p.add_run(f"{row['partner_name']}  (Partner Code: {row['partner_code']})")
    r.bold = True
    last_from_para = p

    if row.get("address"):
        p = doc.add_paragraph(row["address"])
        p.paragraph_format.right_indent = Cm(8.5)
        _tight(p)
        last_from_para = p

    if row.get("pan_no"):
        p = doc.add_paragraph()
        _tight(p)
        r = p.add_run(f"PAN No\t: {row['pan_no']}")
        r.bold = True
        last_from_para = p

    if gst["gst_applicable"]:
        p = doc.add_paragraph()
        _tight(p, space_after=10)
        r = p.add_run(f"GST No\t: {str(row['gst_no']).strip()}")
        r.bold = True
    else:
        last_from_para.paragraph_format.space_after = Pt(10)

    _tight(doc.add_paragraph("To,"))
    p = doc.add_paragraph()
    _tight(p)
    p.add_run(PAYEE_NAME).bold = True
    for line in PAYEE_ADDRESS_LINES:
        _tight(doc.add_paragraph(line))
    p = doc.add_paragraph()
    _tight(p)
    p.add_run(f"GSTIN NO: {PAYEE_GSTIN}").bold = True
    p = doc.add_paragraph()
    _tight(p, space_after=10)
    p.add_run(f"Place of Supply : {PLACE_OF_SUPPLY}").bold = True

    # ---- particulars table (spacious, professional) ------------------------
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Cm(9), Cm(2.5), Cm(4)]
    hdr = table.rows[0].cells
    _set_row_height(table.rows[0], twips=420)
    for i, (text, w) in enumerate(zip(["Particulars", "SAC CODE", "Amount"], widths)):
        hdr[i].width = w
        _shade_cell(hdr[i])
        _set_cell_margins(hdr[i])
        _set_cell_text(hdr[i], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

    def add_row(particulars, sac, amount, bold=False, height=420):
        new_row = table.add_row()
        _set_row_height(new_row, twips=height)
        cells = new_row.cells
        for c, w in zip(cells, widths):
            c.width = w
            _set_cell_margins(c)
        _set_cell_text(cells[0], particulars, bold=bold, size=10.5)
        _set_cell_text(cells[1], sac, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
        _set_cell_text(cells[2], amount, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5)
        return cells

    add_row(
        f"COMMISSION PAYOUT FOR THE MONTH OF : {month_period}",
        SAC_CODE,
        f"Rs. {fmt_rupees(gst['base'])}/-",
        bold=True,
        height=520,
    )
    if gst["gst_applicable"]:
        add_row("IGST @ 18%", "", f"Rs. {fmt_rupees(gst['igst'])}/-" if gst["igst"] else "Rs. 00/-")
        add_row("SGST @ 9%", "", f"Rs. {fmt_rupees(gst['sgst'])}/-" if gst["sgst"] else "Rs. 00/-")
        add_row("CGST @ 9%", "", f"Rs. {fmt_rupees(gst['cgst'])}/-" if gst["cgst"] else "Rs. 00/-")
    add_row("Total", "", f"Rs. {fmt_rupees(gst['total'])}/-", bold=True, height=460)

    # ---- amount / tax in words (two lines, merged across the row) ---------
    words_row = table.add_row()
    _set_row_height(words_row, twips=680)
    words_cells = words_row.cells
    for c in words_cells:
        _set_cell_margins(c)
    merged = words_cells[0].merge(words_cells[1]).merge(words_cells[2])
    merged.text = ""

    p1 = merged.paragraphs[0]
    _tight(p1, space_after=2)
    r1 = p1.add_run(f"Amount Chargeable (in words) : {number_to_indian_words(gst['total'])} RUPEES ONLY")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = "Arial"

    p2 = merged.add_paragraph()
    _tight(p2, space_after=2)
    r2 = p2.add_run(f"Tax Amount (in words) : {number_to_indian_words(gst['tax_total'])} RUPEES ONLY")
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.name = "Arial"

    spacer = doc.add_paragraph()
    _tight(spacer, space_after=6)
    p = doc.add_paragraph()
    _tight(p)
    p.add_run("Bank Details:-").bold = True

    bank_table = doc.add_table(rows=0, cols=3)
    bank_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_border_table(bank_table)
    bank_widths = [Cm(6.2), Cm(0.6), Cm(8.5)]

    def bank_line(label, value):
        cells = bank_table.add_row().cells
        for c, w in zip(cells, bank_widths):
            c.width = w
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], ":")
        _set_cell_text(cells[2], str(value), bold=True)
        for c in cells:
            _tight(c.paragraphs[0], space_after=4)

    bank_line("Beneficiary Account Name", row["partner_name"])
    bank_line("Bank Name", row["bank_name"])
    bank_line("Beneficiary Account Number", str(row["account_number"]).lstrip("'`"))
    bank_line("Beneficiary RTGS Code", row["ifsc"])

    spacer2 = doc.add_paragraph()
    _tight(spacer2, space_after=6)

    # "For <partner>" then extra breathing room before the signatory line so
    # there's room to actually sign in between.
    p_for = doc.add_paragraph()
    _tight(p_for, space_after=40)
    p_for.add_run(f"For {row['partner_name']}").bold = True

    p_sig = doc.add_paragraph()
    _tight(p_sig)
    p_sig.add_run("AUTHORISED SIGNATORY").bold = True

    doc.save(out_path)
    return gst


def safe_filename(text):
    return re.sub(r"[^A-Za-z0-9_\-]+", "_", str(text)).strip("_")
